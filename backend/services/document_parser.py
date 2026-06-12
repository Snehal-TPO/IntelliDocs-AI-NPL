from pathlib import Path

import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}


class DocumentParseError(Exception):
    pass


def extract_text(file_path: str, extension: str) -> str:
    extension = extension.lower()
    path = Path(file_path)

    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(f"Unsupported file type: {extension}")

    try:
        if extension == ".pdf":
            return _extract_pdf(path)
        if extension == ".docx":
            return _extract_docx(path)
        if extension == ".txt":
            return _extract_txt(path)
        if extension in {".png", ".jpg", ".jpeg"}:
            return _extract_image(path)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not extract text: {exc}") from exc

    raise DocumentParseError(f"No parser configured for {extension}")


def _extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                parts.append(f"[Page {page_number}]\n{text}")
    return "\n\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def _extract_image(path: Path) -> str:
    try:
        with Image.open(path) as image:
            return pytesseract.image_to_string(image).strip()
    except pytesseract.TesseractNotFoundError as exc:
        raise DocumentParseError(
            "Tesseract OCR is not installed or is not available on PATH."
        ) from exc
